"""
document_extractor.py - Document Text Extraction Module

This module handles extracting text content from PDF and DOCX files.
It provides a unified interface for document processing regardless of format.

Usage:
    from document_extractor import DocumentExtractor

    extractor = DocumentExtractor()
    text = extractor.extract("document.pdf")
    # or
    text = extractor.extract("document.docx")
"""

import os
import logging
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document


class DocumentExtractor:
    """
    Extracts text content from PDF and DOCX files.

    This class is responsible for:
    - Reading PDF files using PyPDF2
    - Reading DOCX files using python-docx
    - Auto-detecting file format based on extension
    - Returning clean text content

    Supported formats:
        - PDF (.pdf): Extracts text from all pages
        - DOCX (.docx): Extracts text from all paragraphs
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}

    def __init__(self):
        """Initialize the document extractor."""
        self.logger = logging.getLogger(__name__)

    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract all text from a PDF file.

        Reads through all pages of the PDF and combines their text content.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text as a single string with pages separated by newlines

        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: If PDF cannot be read or is corrupted
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        self.logger.info(f"Extracting text from PDF: {file_path}")
        text_parts = []

        try:
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            self.logger.debug(f"PDF has {total_pages} pages")

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    self.logger.debug(f"Page {i+1}: extracted {len(page_text)} characters")

            combined_text = '\n'.join(text_parts)
            self.logger.info(f"Total extracted: {len(combined_text)} characters from {total_pages} pages")
            return combined_text

        except Exception as e:
            self.logger.error(f"Failed to extract PDF: {e}")
            raise

    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract all text from a DOCX file.

        Reads through all paragraphs of the Word document and combines their text.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text as a single string with paragraphs separated by newlines

        Raises:
            FileNotFoundError: If file doesn't exist
            Exception: If DOCX cannot be read or is corrupted
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        self.logger.info(f"Extracting text from DOCX: {file_path}")
        text_parts = []

        try:
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            self.logger.debug(f"DOCX has {total_paragraphs} paragraphs")

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Use double newlines to preserve paragraph boundaries for chunking
            combined_text = '\n\n'.join(text_parts)
            self.logger.info(f"Total extracted: {len(combined_text)} characters from {len(text_parts)} paragraphs")
            return combined_text

        except Exception as e:
            self.logger.error(f"Failed to extract DOCX: {e}")
            raise

    def extract(self, file_path: str) -> str:
        """
        Auto-detect file type and extract text.

        This is the main method to use. It automatically detects the file format
        based on the extension and calls the appropriate extraction method.

        Args:
            file_path: Path to the document (PDF or DOCX)

        Returns:
            Extracted text as a single string

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        # Get file extension (lowercase)
        _, ext = os.path.splitext(file_path.lower())

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                f"Supported formats: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Call appropriate extraction method
        if ext == '.pdf':
            return self.extract_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_from_docx(file_path)

    def get_file_info(self, file_path: str) -> dict:
        """
        Get basic information about a document file.

        Args:
            file_path: Path to the document

        Returns:
            Dictionary with filename, extension, size_bytes, exists
        """
        _, ext = os.path.splitext(file_path.lower())
        exists = os.path.exists(file_path)

        info = {
            'filename': os.path.basename(file_path),
            'extension': ext,
            'exists': exists,
            'size_bytes': os.path.getsize(file_path) if exists else 0,
            'supported': ext in self.SUPPORTED_EXTENSIONS
        }

        return info
